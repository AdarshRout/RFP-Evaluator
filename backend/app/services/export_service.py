import io
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xlsxwriter
from app.schemas.models import EvaluationReport


def export_pdf(report: EvaluationReport) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    navy = colors.HexColor("#1A3A5C")
    blue = colors.HexColor("#2B6CB0")
    light = colors.HexColor("#EBF4FF")

    title_style = ParagraphStyle("Title", parent=styles["Title"], textColor=navy, fontSize=20, spaceAfter=6)
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"], textColor=blue, fontSize=13, spaceBefore=14, spaceAfter=6)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=4)
    verdict_color = colors.HexColor("#276749") if "SHORTLIST" in report.recommendation.upper() else \
                    colors.HexColor("#9B2335") if "REJECT" in report.recommendation.upper() else \
                    colors.HexColor("#744210")

    elements = [
        Paragraph("RFP Evaluation Report", title_style),
        Paragraph(f"Vendor: {report.vendor_name}", ParagraphStyle("Sub", parent=styles["Normal"], fontSize=12, textColor=colors.HexColor("#4A5568"), spaceAfter=4)),
        HRFlowable(width="100%", thickness=2, color=blue, spaceAfter=12),
        Paragraph("Score Summary", h2_style),
    ]

    score_data = [
        ["Metric", "Score"],
        ["Weighted Score", f"{report.weighted_score:.1f} / 10"],
        ["Overall Score", f"{report.total_score:.1f} / 10"],
        ["Requirements Evaluated", str(len(report.requirement_scores))],
    ]
    score_table = Table(score_data, colWidths=[8*cm, 8*cm])
    score_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), navy),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("BACKGROUND", (0, 1), (-1, -1), light),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, light]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    elements += [score_table, Spacer(1, 0.4*cm)]

    elements.append(Paragraph("Verdict", h2_style))
    elements.append(Paragraph(report.recommendation, ParagraphStyle("Verdict", parent=styles["Normal"], fontSize=11, textColor=verdict_color, spaceAfter=8)))

    elements.append(Paragraph("Executive Summary", h2_style))
    elements.append(Paragraph(report.executive_summary, body_style))
    elements.append(Spacer(1, 0.3*cm))

    elements.append(Paragraph("Category Scores", h2_style))
    cat_data = [["Category", "Score"]] + [[k, f"{v:.1f} / 10"] for k, v in report.category_scores.items()]
    cat_table = Table(cat_data, colWidths=[8*cm, 8*cm])
    cat_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), blue),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, light]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    elements += [cat_table, Spacer(1, 0.3*cm)]

    if report.strengths:
        elements.append(Paragraph("Strengths", h2_style))
        for s in report.strengths:
            elements.append(Paragraph(f"• {s}", body_style))
        elements.append(Spacer(1, 0.2*cm))

    if report.gaps:
        elements.append(Paragraph("Gaps & Risks", h2_style))
        for g in report.gaps:
            elements.append(Paragraph(f"• {g}", body_style))
        elements.append(Spacer(1, 0.3*cm))

    elements.append(Paragraph("Requirement-Level Breakdown", h2_style))
    table_cell_style = ParagraphStyle("TableCell", parent=styles["Normal"], fontSize=9, leading=12)
    table_hdr_style = ParagraphStyle("TableHdr", parent=styles["Normal"], fontSize=9, leading=12, textColor=colors.white, fontName="Helvetica-Bold")

    req_data = [[
        Paragraph("Req ID", table_hdr_style),
        Paragraph("Score", table_hdr_style),
        Paragraph("Justification", table_hdr_style)
    ]]
    for ps in report.requirement_scores:
        req_data.append([
            Paragraph(ps.requirement_id, table_cell_style),
            Paragraph(f"{ps.score:.1f}", table_cell_style),
            Paragraph(ps.justification, table_cell_style)
        ])
    req_table = Table(req_data, colWidths=[2.5*cm, 2*cm, 12.5*cm])
    req_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), navy),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#DDDDDD")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, light]),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    elements.append(req_table)

    doc.build(elements)
    return buf.getvalue()


def export_docx(report: EvaluationReport) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    def _heading(text: str, level: int = 1):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C) if level == 1 else RGBColor(0x2B, 0x6C, 0xB0)

    _heading("RFP Evaluation Report")
    doc.add_paragraph(f"Vendor: {report.vendor_name}").runs[0].bold = True

    _heading("Score Summary", 2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text, hdr[1].text = "Metric", "Score"
    for metric, value in [("Weighted Score", f"{report.weighted_score:.1f}/10"),
                           ("Overall Score", f"{report.total_score:.1f}/10"),
                           ("Requirements", str(len(report.requirement_scores)))]:
        row = tbl.add_row().cells
        row[0].text, row[1].text = metric, value

    _heading("Verdict", 2)
    doc.add_paragraph(report.recommendation)

    _heading("Executive Summary", 2)
    doc.add_paragraph(report.executive_summary)

    _heading("Category Scores", 2)
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.style = "Table Grid"
    tbl2.rows[0].cells[0].text, tbl2.rows[0].cells[1].text = "Category", "Score"
    for cat, score in report.category_scores.items():
        row = tbl2.add_row().cells
        row[0].text, row[1].text = cat, f"{score:.1f}/10"

    _heading("Strengths", 2)
    for s in report.strengths:
        doc.add_paragraph(s, style="List Bullet")

    _heading("Gaps & Risks", 2)
    for g in report.gaps:
        doc.add_paragraph(g, style="List Bullet")

    _heading("Requirement Breakdown", 2)
    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.style = "Table Grid"
    hdrs = ["Req ID", "Score", "Justification", "Evidence"]
    for i, h in enumerate(hdrs):
        tbl3.rows[0].cells[i].text = h
    for ps in report.requirement_scores:
        row = tbl3.add_row().cells
        row[0].text = ps.requirement_id
        row[1].text = f"{ps.score:.1f}"
        row[2].text = ps.justification
        row[3].text = ps.evidence[:100]

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_xlsx(report: EvaluationReport) -> bytes:
    buf = io.BytesIO()
    wb = xlsxwriter.Workbook(buf, {"in_memory": True})

    navy = "#1A3A5C"
    blue = "#2B6CB0"
    white = "#FFFFFF"
    light = "#EBF4FF"

    hdr_fmt = wb.add_format({"bold": True, "bg_color": navy, "font_color": white, "border": 1, "font_size": 11})
    sub_hdr_fmt = wb.add_format({"bold": True, "bg_color": blue, "font_color": white, "border": 1})
    cell_fmt = wb.add_format({"border": 1, "font_size": 10, "text_wrap": True, "valign": "top"})
    alt_fmt = wb.add_format({"border": 1, "bg_color": light, "font_size": 10, "text_wrap": True, "valign": "top"})
    score_fmt = wb.add_format({"border": 1, "bold": True, "font_size": 11, "num_format": "0.0"})

    ws1 = wb.add_worksheet("Summary")
    ws1.set_column("A:A", 28)
    ws1.set_column("B:B", 20)
    ws1.write("A1", "RFP Evaluation Report", wb.add_format({"bold": True, "font_size": 16, "font_color": navy}))
    ws1.write("A2", f"Vendor: {report.vendor_name}", wb.add_format({"font_size": 12, "font_color": blue}))
    ws1.write("A4", "Metric", hdr_fmt)
    ws1.write("B4", "Score", hdr_fmt)
    for i, (k, v) in enumerate([("Weighted Score", f"{report.weighted_score:.1f}/10"),
                                  ("Overall Score", f"{report.total_score:.1f}/10"),
                                  ("Requirements Evaluated", str(len(report.requirement_scores))),
                                  ("Verdict", report.recommendation)]):
        fmt = alt_fmt if i % 2 else cell_fmt
        ws1.write(4+i, 0, k, fmt)
        ws1.write(4+i, 1, v, fmt)

    ws1.write("A10", "Executive Summary", sub_hdr_fmt)
    ws1.merge_range("A11:B13", report.executive_summary, wb.add_format({"text_wrap": True, "border": 1, "font_size": 10, "valign": "top"}))
    ws1.set_row(10, 60)

    ws2 = wb.add_worksheet("Category Scores")
    ws2.set_column("A:A", 22)
    ws2.set_column("B:B", 15)
    ws2.write(0, 0, "Category", hdr_fmt)
    ws2.write(0, 1, "Score (/10)", hdr_fmt)
    for i, (cat, score) in enumerate(report.category_scores.items()):
        fmt = alt_fmt if i % 2 else cell_fmt
        ws2.write(i+1, 0, cat, fmt)
        ws2.write(i+1, 1, score, score_fmt)

    ws3 = wb.add_worksheet("Requirement Scores")
    ws3.set_column("A:A", 12)
    ws3.set_column("B:B", 10)
    ws3.set_column("C:C", 40)
    ws3.set_column("D:D", 40)
    headers = ["Req ID", "Score", "Justification", "Evidence"]
    for j, h in enumerate(headers):
        ws3.write(0, j, h, hdr_fmt)
    for i, ps in enumerate(report.requirement_scores):
        fmt = alt_fmt if i % 2 else cell_fmt
        ws3.write(i+1, 0, ps.requirement_id, fmt)
        ws3.write(i+1, 1, ps.score, score_fmt)
        ws3.write(i+1, 2, ps.justification, fmt)
        ws3.write(i+1, 3, ps.evidence, fmt)
        ws3.set_row(i+1, 45)

    ws4 = wb.add_worksheet("Strengths & Gaps")
    ws4.set_column("A:A", 15)
    ws4.set_column("B:B", 70)
    ws4.write(0, 0, "Type", hdr_fmt)
    ws4.write(0, 1, "Detail", hdr_fmt)
    row = 1
    for s in report.strengths:
        ws4.write(row, 0, "Strength", wb.add_format({"bg_color": "#F0FFF4", "font_color": "#276749", "border": 1, "bold": True}))
        ws4.write(row, 1, s, cell_fmt)
        row += 1
    for g in report.gaps:
        ws4.write(row, 0, "Gap", wb.add_format({"bg_color": "#FFF5F5", "font_color": "#9B2335", "border": 1, "bold": True}))
        ws4.write(row, 1, g, cell_fmt)
        row += 1

    wb.close()
    return buf.getvalue()
